# app.py
import io
import re
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

st.set_page_config(page_title="Estrattore evidenziati e commenti .docx", layout="wide")

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# -----------------------------
# Utilities
# -----------------------------

def normalize_highlight_value(val):
    """Normalize python-docx WD_COLOR_INDEX to lowercase string name or None."""
    if val is None:
        return None
    s = str(val)
    name = s.split(".")[-1] if "." in s else s
    return name.strip().lower()


def iter_paragraphs(doc):
    # regular paragraphs
    for p in doc.paragraphs:
        yield p
    # paragraphs in tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                seen = set()
                for p in cell.paragraphs:
                    pid = id(p._element)
                    if pid in seen:
                        continue
                    seen.add(pid)
                    yield p


def sentence_window(text: str, start_index: int) -> str:
    parts = re.split(r"(?<=[\.!?])\s+", text)
    cum = 0
    for s in parts:
        if cum <= start_index < cum + len(s) + 1:
            return s.strip()
        cum += len(s) + 1
    return text.strip()


# -----------------------------
# Extraction
# -----------------------------

def extract_highlights(doc, filename):
    """Merge contiguous runs with the same highlight color within a paragraph."""
    rows = []
    for p in iter_paragraphs(doc):
        para_text = p.text or ""
        if not para_text.strip():
            continue
        idx = 0
        buf = []
        cur_color = None
        seg_start = None

        def flush():
            nonlocal buf, cur_color, seg_start
            if cur_color is None or not buf:
                buf, cur_color, seg_start = [], None, None
                return
            text_joined = "".join(buf).strip()
            if text_joined:
                rows.append(
                    {
                        "filename": filename,
                        "type": "highlight",
                        "highlight_color": cur_color,
                        "extracted_text": re.sub(r"\s+", " ", text_joined),
                        "context": sentence_window(para_text, seg_start or 0),
                        "paragraph": para_text.strip(),
                    }
                )
            buf, cur_color, seg_start = [], None, None

        for r in p.runs:
            text = r.text or ""
            rlen = len(text)
            color = normalize_highlight_value(r.font.highlight_color)
            if rlen:
                if color is not None:
                    if cur_color is None:
                        cur_color = color
                        seg_start = idx
                        buf = [text]
                    elif color == cur_color:
                        buf.append(text)
                    else:
                        flush()
                        cur_color = color
                        seg_start = idx
                        buf = [text]
                else:
                    flush()
                idx += rlen
        flush()
    return rows


def get_comments_map(doc):
    comments = {}
    part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            part = rel.target_part
            break
    if part is None:
        return comments
    root = etree.fromstring(part.blob)
    for c in root.xpath(".//w:comment", namespaces=NS):
        cid = int(c.get("{%s}id" % NS["w"]))
        author = c.get("{%s}author" % NS["w"]) or ""
        date = c.get("{%s}date" % NS["w"]) or ""
        text_runs = c.xpath(".//w:t", namespaces=NS)
        ctext = "".join([t.text or "" for t in text_runs]).strip()
        comments[cid] = {"author": author, "date": date, "comment_text": ctext}
    return comments


def get_commented_spans(doc):
    body = doc.element.body
    open_ids = []
    spans = {}

    def walk(el):
        nonlocal open_ids
        if el.tag == "{%s}commentRangeStart" % NS["w"]:
            cid = int(el.get("{%s}id" % NS["w"]))
            open_ids.append(cid)
            spans.setdefault(cid, [])
        elif el.tag == "{%s}commentRangeEnd" % NS["w"]:
            cid = int(el.get("{%s}id" % NS["w"]))
            if cid in open_ids:
                open_ids.remove(cid)
        if el.tag == "{%s}t" % NS["w"] and el.text:
            for cid in open_ids:
                spans.setdefault(cid, []).append(el.text)
        for child in el:
            walk(child)

    walk(body)
    return {cid: "".join(parts).strip() for cid, parts in spans.items() if "".join(parts).strip()}


def parse_code_labels(text):
    results = []
    if not text:
        return results
    t = text.strip()
    # XML-like pairs
    xml_pairs = re.findall(
        r"<\s*codice\s*>(.*?)<\s*/?\\?\s*codice\s*>\s*<\s*commento\s*>(.*?)<\s*/?\\?\s*commento\s*>",
        t,
        flags=re.IGNORECASE | re.DOTALL,
    )
    for code, label in xml_pairs:
        code = (code or "").strip() or None
        label = (label or "").strip() or None
        if code or label:
            results.append((code, label))
    # key=value pairs
    kv_pairs = re.findall(
        r"(?:codice|code)\s*=\s*([A-Za-z0-9_]+)\s*(?:;|,|\||\n|\r|\s)+(?:commento|label)\s*=\s*([^;,\|\n\r]+)",
        t,
        flags=re.IGNORECASE,
    )
    for code, label in kv_pairs:
        code = code.strip()
        label = label.strip()
        if code or label:
            results.append((code or None, label or None))
    # line-based patterns
    normalized = re.sub(r"[;|]+", "\n", t)
    for line in normalized.splitlines():
        s = line.strip()
        if not s:
            continue
        m_br = re.match(r"\[\s*([A-Za-z0-9_]+)\s*\]\s+(.+)$", s)
        if m_br:
            results.append((m_br.group(1).strip(), m_br.group(2).strip()))
            continue
        m_col = re.match(r"([A-Za-z0-9_]+)\s*[:\-–]\s*(.+)$", s)
        if m_col:
            results.append((m_col.group(1).strip(), m_col.group(2).strip()))
            continue
    # dedupe
    seen = set()
    deduped = []
    for code, label in results:
        key = (code or "", label or "")
        if key not in seen and (code or label):
            seen.add(key)
            deduped.append((code, label))
    return deduped


def extract_comments(doc, filename):
    comments_meta = get_comments_map(doc)
    quoted_map = get_commented_spans(doc)
    rows = []
    for cid, meta in comments_meta.items():
        raw = meta.get("comment_text", "")
        pairs = parse_code_labels(raw)
        if not pairs:
            rows.append(
                {
                    "filename": filename,
                    "type": "comment",
                    "comment_id": cid,
                    "author": meta.get("author", ""),
                    "date": meta.get("date", ""),
                    "quoted_text": quoted_map.get(cid, ""),
                    "comment_text": raw,
                    "code": None,
                    "label": None,
                }
            )
        else:
            for code, label in pairs:
                rows.append(
                    {
                        "filename": filename,
                        "type": "comment",
                        "comment_id": cid,
                        "author": meta.get("author", ""),
                        "date": meta.get("date", ""),
                        "quoted_text": quoted_map.get(cid, ""),
                        "comment_text": raw,
                        "code": code,
                        "label": label,
                    }
                )
    return rows


def process_docx_bytes(file_bytes, filename):
    doc = Document(io.BytesIO(file_bytes))
    rows = []
    rows += extract_highlights(doc, filename)
    rows += extract_comments(doc, filename)
    return rows


# -----------------------------
# Export helpers
# -----------------------------

def export_excel(high_df, com_df, link_df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if high_df is not None and not high_df.empty:
            high_order = [
                "filename",
                "extracted_text",
                "highlight_color",
                "context",
                "paragraph",
            ]
            for col in high_order:
                if col not in high_df.columns:
                    high_df[col] = ""
            high_df[high_order].to_excel(writer, index=False, sheet_name="Highlights")
        if com_df is not None and not com_df.empty:
            com_order = [
                "filename",
                "comment_id",
                "author",
                "date",
                "quoted_text",
                "comment_text",
                "code",
                "label",
            ]
            for col in com_order:
                if col not in com_df.columns:
                    com_df[col] = ""
            com_df[com_order].to_excel(
                writer, index=False, sheet_name="Commenti (esplosi)"
            )
        if link_df is not None and not link_df.empty:
            link_order = [
                "filename",
                "comment_id",
                "author",
                "date",
                "code",
                "label",
                "quoted_text",
                "highlight_matches",
                "highlights_concat",
            ]
            for col in link_order:
                if col not in link_df.columns:
                    link_df[col] = ""
            link_df[link_order].to_excel(
                writer, index=False, sheet_name="Annotazioni collegate"
            )
        summary = pd.DataFrame(
            {
                "sheet": [
                    "Highlights",
                    "Commenti (esplosi)",
                    "Annotazioni collegate",
                ],
                "rows": [
                    0 if high_df is None else len(high_df),
                    0 if com_df is None else len(com_df),
                    0 if link_df is None else len(link_df),
                ],
            }
        )
        summary.to_excel(writer, index=False, sheet_name="Riepilogo")
    output.seek(0)
    return output


# -----------------------------
# UI renderer
# -----------------------------

def render_app_views(df: pd.DataFrame):
    # Sidebar filters
    st.sidebar.header("Filtri")
    filenames = sorted(df["filename"].dropna().unique().tolist())
    sel_files = st.sidebar.multiselect("File", filenames, default=filenames)

    types = sorted(df["type"].dropna().unique().tolist())
    sel_types = st.sidebar.multiselect("Tipo annotazione", types, default=types)

    colors = (
        sorted(df.loc[df["type"] == "highlight", "highlight_color"].dropna().unique().tolist())
        if "highlight_color" in df.columns
        else []
    )
    sel_colors = st.sidebar.multiselect(
        "Colore evidenziazione", colors, default=colors
    )

    authors = (
        sorted(df.loc[df["type"] == "comment", "author"].dropna().unique().tolist())
        if "author" in df.columns
        else []
    )
    sel_authors = st.sidebar.multiselect(
        "Autore commento", authors, default=authors
    )

    codes = (
        sorted(df.loc[df["type"] == "comment", "code"].dropna().unique().tolist())
        if "code" in df.columns
        else []
    )
    sel_codes = st.sidebar.multiselect("Codice (dal commento)", codes, default=codes)

    labels = (
        sorted(df.loc[df["type"] == "comment", "label"].dropna().unique().tolist())
        if "label" in df.columns
        else []
    )
    sel_labels = st.sidebar.multiselect("Etichetta (dal commento)", labels, default=labels)

    doc_types = (
        sorted(df["doc_type"].dropna().unique().tolist()) if "doc_type" in df.columns else []
    )
    sel_doc_types = st.sidebar.multiselect(
        "Tipo documento", doc_types, default=doc_types
    )

    search_text = st.sidebar.text_input("Cerca testo (qualsiasi campo)")
    date_min = st.sidebar.text_input("Data minima (YYYY-MM-DD)")
    date_max = st.sidebar.text_input("Data massima (YYYY-MM-DD)")

    # Apply filters
    fdf = df.copy()
    if sel_files:
        fdf = fdf[fdf["filename"].isin(sel_files)]
    if sel_types:
        fdf = fdf[fdf["type"].isin(sel_types)]
    if sel_colors:
        fdf = fdf[(fdf["type"] != "highlight") | (fdf["highlight_color"].isin(sel_colors))]
    if sel_doc_types:
        fdf = fdf[fdf["doc_type"].isin(sel_doc_types)]
    if sel_authors:
        fdf = fdf[(fdf["type"] != "comment") | (fdf["author"].isin(sel_authors))]
    if sel_codes:
        fdf = fdf[(fdf["type"] != "comment") | (fdf["code"].isin(sel_codes))]
    if sel_labels:
        fdf = fdf[(fdf["type"] != "comment") | (fdf["label"].isin(sel_labels))]

    if search_text:
        mask = pd.Series(False, index=fdf.index)
        for col in [
            "extracted_text",
            "paragraph",
            "context",
            "comment_text",
            "quoted_text",
            "author",
            "filename",
            "code",
            "label",
        ]:
            if col in fdf.columns:
                mask = mask | fdf[col].astype(str).str.contains(
                    search_text, case=False, na=False
                )
        fdf = fdf[mask]

    def parse_date(val):
        if not val:
            return None
        try:
            return datetime.fromisoformat(val.replace("Z", "").replace("+00:00", ""))
        except Exception:
            for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S"):
                try:
                    return datetime.strptime(val[: len(fmt)], fmt)
                except Exception:
                    pass
        return None

    dmin = parse_date(date_min) if date_min else None
    dmax = parse_date(date_max) if date_max else None
    if dmin or dmax:
        if "date" in fdf.columns:
            dt_series = fdf["date"].apply(parse_date)
            if dmin:
                fdf = fdf[(dt_series.isna()) | (dt_series >= dmin)]
            if dmax:
                fdf = fdf[(dt_series.isna()) | (dt_series <= dmax)]

    high_df = fdf[fdf["type"] == "highlight"].copy()
    com_df = fdf[fdf["type"] == "comment"].copy()

    # Scrematura (simplified view)
    st.subheader("Scrematura (evidenziati) – vista semplificata")
    screm_df = pd.DataFrame()
    if not high_df.empty:
        cols = [
            c
            for c in [
                "extracted_text",
                "highlight_color",
                "macro_category",
                "subject",
                "filename",
                "doc_type",
            ]
            if c in high_df.columns
        ]
        screm_df = high_df[cols].rename(
            columns={
                "extracted_text": "testo_evidenziato",
                "highlight_color": "colore",
                "macro_category": "categoria",
                "subject": "soggetto",
                "filename": "file",
                "doc_type": "tipo_documento",
            }
        )

        group_toggle = st.checkbox(
            "Raggruppa per file + categoria (concatena testi)",
            value=False,
            key="grp_screm",
        )
        if group_toggle:
            gcols = [c for c in ["file", "tipo_documento", "categoria"] if c in screm_df.columns]
            if gcols:
                grouped = (
                    screm_df.groupby(gcols, dropna=False)["testo_evidenziato"]
                    .apply(lambda s: " | ".join([t for t in s.astype(str).tolist() if t]))
                    .reset_index()
                )
                st.dataframe(grouped, use_container_width=True, hide_index=True)
            else:
                st.dataframe(screm_df, use_container_width=True, hide_index=True)
        else:
            st.dataframe(screm_df, use_container_width=True, hide_index=True)
    else:
        st.info("Nessun evidenziato dopo i filtri per la vista semplificata.")

    # Linked annotations view
    link_rows = []
    if not com_df.empty:
        all_high = df[df["type"] == "highlight"].copy()
        for _, row in com_df.iterrows():
            quoted = str(row.get("quoted_text") or "").strip()
            fname = row.get("filename")
            matches = []
            if quoted:
                cand = all_high[(all_high["filename"] == fname)]
                for _, hr in cand.iterrows():
                    htext = str(hr.get("extracted_text") or "").strip()
                    if not htext:
                        continue
                    qn = re.sub(r"\s+", " ", quoted).strip()
                    hn = re.sub(r"\s+", " ", htext).strip()
                    if hn and qn and (hn in qn or qn in hn):
                        matches.append(htext)
            link_rows.append(
                {
                    "filename": fname,
                    "comment_id": row.get("comment_id"),
                    "author": row.get("author"),
                    "date": row.get("date"),
                    "code": row.get("code"),
                    "label": row.get("label"),
                    "quoted_text": quoted,
                    "highlight_matches": len(matches),
                    "highlights_concat": " | ".join(sorted(set(matches))) if matches else "",
                }
            )
    link_df = pd.DataFrame(link_rows) if link_rows else pd.DataFrame()

    st.subheader("Annotazioni collegate (Commento ↔ Evidenziati)")
    if not link_df.empty:
        default_link_cols = [
            "filename",
            "comment_id",
            "author",
            "date",
            "code",
            "label",
            "quoted_text",
            "highlight_matches",
            "highlights_concat",
        ]
        link_cols = [c for c in default_link_cols if c in link_df.columns]
        sel_link_cols = st.multiselect(
            "Colonne da mostrare (annotazioni collegate)",
            link_df.columns.tolist(),
            default=link_cols,
            key="link_cols",
        )
        st.dataframe(link_df[sel_link_cols], use_container_width=True, hide_index=True)
    else:
        st.info("Nessuna annotazione collegata dopo i filtri.")

    st.subheader("Evidenziati")
    if not high_df.empty:
        default_high_cols = [
            "filename",
            "extracted_text",
            "highlight_color",
            "context",
            "paragraph",
        ]
        high_cols = [c for c in default_high_cols if c in high_df.columns]
        sel_high_cols = st.multiselect(
            "Colonne da mostrare (evidenziati)",
            high_df.columns.tolist(),
            default=high_cols,
            key="high_cols",
        )
        st.dataframe(high_df[sel_high_cols], use_container_width=True, hide_index=True)
    else:
        st.info("Nessun evidenziato dopo i filtri.")

    st.subheader("Commenti (esplosi per codice/etichetta)")
    if not com_df.empty:
        default_com_cols = [
            "filename",
            "comment_id",
            "author",
            "date",
            "quoted_text",
            "comment_text",
            "code",
            "label",
        ]
        com_cols = [c for c in default_com_cols if c in com_df.columns]
        sel_com_cols = st.multiselect(
            "Colonne da mostrare (commenti)",
            com_df.columns.tolist(),
            default=com_cols,
            key="com_cols",
        )
        st.dataframe(com_df[sel_com_cols], use_container_width=True, hide_index=True)
    else:
        st.info("Nessun commento dopo i filtri.")

    st.divider()
    st.subheader("Esportazione")
    excel_bytes = export_excel(high_df, com_df, link_df)
    st.download_button(
        label="Scarica Excel filtrato",
        data=excel_bytes,
        file_name="estrazioni_filtrate.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    # Exports for Interviste and Focus group from scrematura
    if not screm_df.empty and "tipo_documento" in screm_df.columns:
        int_df = screm_df[screm_df["tipo_documento"] == "Intervista"].copy()
        fg_df = screm_df[screm_df["tipo_documento"] == "Focus group"].copy()

        def to_excel_bytes(df_in):
            out = io.BytesIO()
            with pd.ExcelWriter(out, engine="openpyxl") as w:
                df_in.to_excel(w, index=False, sheet_name="Scrematura")
            out.seek(0)
            return out

        if not int_df.empty:
            st.download_button(
                label="Scarica Interviste (XLSX)",
                data=to_excel_bytes(int_df),
                file_name="interviste_scrematura.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            st.download_button(
                label="Scarica Interviste (CSV)",
                data=int_df.to_csv(index=False).encode("utf-8"),
                file_name="interviste_scrematura.csv",
                mime="text/csv",
            )
        if not fg_df.empty:
            st.download_button(
                label="Scarica Focus group (XLSX)",
                data=to_excel_bytes(fg_df),
                file_name="focus_group_scrematura.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            st.download_button(
                label="Scarica Focus group (CSV)",
                data=fg_df.to_csv(index=False).encode("utf-8"),
                file_name="focus_group_scrematura.csv",
                mime="text/csv",
            )


# -----------------------------
# Main app
# -----------------------------

st.title("Estrattore evidenziati e commenti da Word (.docx)")
st.caption(
    "Supporto a unione evidenziati per colore, metadati per file e export separati."
)

uploaded_files = st.file_uploader(
    "Carica file .docx", accept_multiple_files=True, type=["docx"]
)

if uploaded_files:
    # Sidebar metadata
    st.sidebar.header("Metadati per file")
    if "file_meta" not in st.session_state:
        st.session_state["file_meta"] = {}
    for up in uploaded_files:
        meta = st.session_state["file_meta"].get(up.name, {})
        with st.sidebar.expander(f"{up.name}"):
            subject = st.text_input(
                "Soggetto intervistato", value=meta.get("subject", ""), key=f"subj_{up.name}"
            )
            doc_type = st.selectbox(
                "Tipo documento",
                ["Intervista", "Focus group"],
                index=0 if meta.get("doc_type", "Intervista") == "Intervista" else 1,
                key=f"dtype_{up.name}",
            )
        st.session_state["file_meta"][up.name] = {
            "subject": subject,
            "doc_type": doc_type,
        }

    # Extract
    rows = []
    for up in uploaded_files:
        try:
            file_bytes = up.getbuffer()
            rows.extend(process_docx_bytes(file_bytes, up.name))
        except Exception as e:
            st.error(f"Errore nel file {up.name}: {e}")

    if not rows:
        st.warning("Nessuna evidenziazione o commento trovati nei documenti caricati.")
    else:
        df = pd.DataFrame(rows)

        # Attach metadata
        file_meta_map_subject = {
            fn: md.get("subject", "") for fn, md in st.session_state.get("file_meta", {}).items()
        }
        file_meta_map_dtype = {
            fn: md.get("doc_type", "Intervista")
            for fn, md in st.session_state.get("file_meta", {}).items()
        }
        if "filename" in df.columns:
            df["subject"] = df["filename"].map(file_meta_map_subject).fillna("")
            df["doc_type"] = df["filename"].map(file_meta_map_dtype).fillna("Intervista")

        # Normalize colors and build mapping
        if "highlight_color" in df.columns:
            df.loc[df["type"] == "highlight", "highlight_color"] = (
                df.loc[df["type"] == "highlight", "highlight_color"].astype(str).str.lower()
            )
            present_colors = sorted(
                [
                    c
                    for c in df.loc[
                        df["type"] == "highlight", "highlight_color"
                    ]
                    .dropna()
                    .unique()
                    .tolist()
                    if c and c != "none"
                ]
            )
        else:
            present_colors = []

        default_color_map = {c: c for c in present_colors}
        if "color_map" not in st.session_state:
            st.session_state["color_map"] = default_color_map.copy()
        else:
            for c in present_colors:
                st.session_state["color_map"].setdefault(c, c)

        with st.sidebar.expander("Mappa Colore → Categoria"):
            updated_map = {}
            for c in present_colors:
                updated_map[c] = st.text_input(
                    f"Categoria per '{c}'",
                    value=st.session_state["color_map"].get(c, c),
                    key=f"cat_{c}",
                )
            st.session_state["color_map"].update(updated_map)

        # Apply macro-category to highlights
        if present_colors:
            df["macro_category"] = None
            mask_high = df["type"] == "highlight"
            df.loc[mask_high, "macro_category"] = (
                df.loc[mask_high, "highlight_color"]
                .map(st.session_state["color_map"])
                .fillna("")
            )

        # Render app views
        render_app_views(df)
else:
    st.info("Carica uno o più file .docx per iniziare.")
