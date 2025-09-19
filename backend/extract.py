import io
import re
from typing import List, Dict, Any, Tuple

import pandas as pd
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree
from docx.oxml.ns import qn

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
VALID_PREFIXES = {"CE", "CS", "BE", "IN", "CC", "A"}

# Map common shading hex codes to the same identifiers returned by python-docx highlights
HEX_TO_HIGHLIGHT = {
    "ffff00": "yellow",
    "00ff00": "bright_green",
    "00ffff": "turquoise",
    "ff00ff": "pink",
    "0000ff": "blue",
    "ff0000": "red",
    "000080": "dark_blue",
    "008080": "teal",
    "008000": "green",
    "800080": "violet",
    "800000": "dark_red",
    "808000": "dark_yellow",
    "7f7f7f": "gray_50",
    "bfbfbf": "gray_25",
    "000000": "black",
    "ffffff": "white",
}


def normalize_highlight_value(val):
    if val is None:
        return None
    s = str(val).strip()
    if not s:
        return None
    if s.startswith("#"):
        s = s[1:]
    name = s.split(".")[-1] if "." in s else s
    lowered = name.strip().lower()
    if lowered in HEX_TO_HIGHLIGHT:
        return HEX_TO_HIGHLIGHT[lowered]
    # fallback: keep hex strings normalized without trailing alpha channel
    if re.fullmatch(r"[0-9a-f]{8}", lowered):
        lowered = lowered[:6]
    return lowered


def run_highlight_or_shading(run) -> str | None:
    """Return normalized highlight color, supporting both highlight and shading."""
    color = normalize_highlight_value(run.font.highlight_color)
    if color:
        return color
    # Inspect shading (<w:shd>) which some docs use instead of highlight
    shd_elems = run._element.xpath("w:rPr/w:shd")
    if not shd_elems:
        return None
    shd = shd_elems[0]
    fill = shd.get(qn("w:fill"))
    if fill and fill.lower() != "auto":
        normalized = normalize_highlight_value(fill)
        if normalized:
            return normalized
    color_attr = shd.get(qn("w:color"))
    if color_attr and color_attr.lower() not in {"auto", "000000"}:
        normalized = normalize_highlight_value(color_attr)
        if normalized:
            return normalized
    theme_fill = shd.get(qn("w:themeFill"))
    if theme_fill:
        return normalize_highlight_value(theme_fill)
    theme_color = shd.get(qn("w:themeColor"))
    if theme_color:
        return normalize_highlight_value(theme_color)
    return None


def iter_paragraphs_with_index(doc: Document):
    idx = 0
    for i, p in enumerate(doc.paragraphs):
        yield p, i
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                seen = set()
                for p in cell.paragraphs:
                    pid = id(p._element)
                    if pid in seen:
                        continue
                    seen.add(pid)
                    # paragraph indices in tables are appended at the end
                    yield p, None


def sentence_window(text: str, start_index: int) -> str:
    parts = re.split(r"(?<=[\.!?])\s+", text)
    cum = 0
    for s in parts:
        if cum <= start_index < cum + len(s) + 1:
            return s.strip()
        cum += len(s) + 1
    return text.strip()


def extract_highlights_with_offsets(doc: Document, filename: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    rows: List[Dict[str, Any]] = []
    paragraphs: List[Dict[str, Any]] = []
    for p, para_index in iter_paragraphs_with_index(doc):
        para_text = p.text or ""
        paragraphs.append(
            {
                "filename": filename,
                "para_index": para_index,
                "text": para_text,
            }
        )
        if not para_text.strip():
            continue
        idx = 0
        buf: List[str] = []
        cur_color: Any = None
        seg_start: Any = None

        def do_flush():
            if cur_color is None or not buf:
                return
            text_joined = "".join(buf)
            if text_joined.strip():
                start = seg_start or 0
                end = start + len(text_joined)
                rows.append(
                    {
                        "filename": filename,
                        "type": "highlight",
                        "highlight_color": cur_color,
                        "text": text_joined,
                        "context": sentence_window(para_text, start),
                        "paragraph": para_text,
                        "offset_start": start,
                        "offset_end": end,
                        "para_index": para_index,
                    }
                )

        for r in p.runs:
            text = r.text or ""
            rlen = len(text)
            color = run_highlight_or_shading(r)
            if rlen:
                if color is not None:
                    if cur_color is None:
                        cur_color = color
                        seg_start = idx
                        buf = [text]
                    elif color == cur_color:
                        buf.append(text)
                    else:
                        do_flush()
                        cur_color = color
                        seg_start = idx
                        buf = [text]
                else:
                    # end of a highlighted segment
                    do_flush()
                    buf = []
                    cur_color = None
                    seg_start = None
                idx += rlen
        # flush any remaining buffered segment
        do_flush()
        buf = []
        cur_color = None
        seg_start = None
    return rows, paragraphs


def get_comments_map(doc: Document):
    comments = {}
    part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            part = rel.target_part
            break
    if part is None:
        return comments
    root = etree.fromstring(part.blob)
    for c in root.xpath(".//w:comment", namespaces={"w": NS["w"]}):
        cid = int(c.get("{%s}id" % NS["w"]))
        author = c.get("{%s}author" % NS["w"]) or ""
        date = c.get("{%s}date" % NS["w"]) or ""
        text_runs = c.xpath(".//w:t", namespaces={"w": NS["w"]})
        ctext = "".join([t.text or "" for t in text_runs]).strip()
        comments[cid] = {"author": author, "date": date, "comment": ctext}
    return comments


def get_commented_spans(doc: Document):
    body = doc.element.body
    open_ids = []
    spans = {}

    def walk(el):
        nonlocal open_ids
        if el.tag == "{%s}commentRangeStart" % NS["w"]:
            cid = int(el.get("{%s}id" % NS["w"]))
            open_ids.append(cid)
            spans.setdefault(cid, [])
        elif el.tag == "{%s}commentRangeEnd" % NS["w"]:
            cid = int(el.get("{%s}id" % NS["w"]))
            if cid in open_ids:
                open_ids.remove(cid)
        if el.tag == "{%s}t" % NS["w"] and el.text:
            for cid in open_ids:
                spans.setdefault(cid, []).append(el.text)
        for child in el:
            walk(child)

    walk(body)
    return {cid: "".join(parts).strip() for cid, parts in spans.items() if "".join(parts).strip()}


def parse_docx(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    doc = Document(io.BytesIO(file_bytes))
    highlights, paragraphs = extract_highlights_with_offsets(doc, filename)
    comments_map = get_comments_map(doc)
    quoted_map = get_commented_spans(doc)

    comments_list: List[Dict[str, Any]] = []
    code_re = re.compile(r"\b[A-Z]{1,4}(?:_[A-Z]{1,4})?\b")
    for cid, meta in comments_map.items():
        text = meta.get("comment", "")
        raw_codes = code_re.findall(text or "") if text else []
        codes: List[str] = []
        for token in raw_codes:
            normalized = (token or "").strip().upper()
            if not normalized:
                continue
            prefix = normalized.split("_")[0]
            if prefix not in VALID_PREFIXES:
                continue
            if normalized in codes:
                continue
            codes.append(normalized)
        code = codes[0] if codes else None
        comments_list.append(
            {
                "filename": filename,
                "id": cid,
                "author": meta.get("author", ""),
                "date": meta.get("date", ""),
                "text": text,
                "quoted": quoted_map.get(cid, ""),
                "code": code,
                "codes": codes,
            }
        )
    return {"highlights": highlights, "comments": comments_list, "paragraphs": paragraphs}
